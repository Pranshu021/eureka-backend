import re, os, tempfile
from app.nodes.prompts import SENIOR_ANALYST_PROMPT
from langchain_core.messages import HumanMessage
from app.utils import logger
from docx import Document
from dotenv import load_dotenv
load_dotenv()
from langsmith import traceable
from app.config import anthropic_llm, gemini_llm

@traceable(run_type="chain", name="Senior Analyst Node")
async def senior_analyst(state):
    junior_analyst_one_report = state.get("junior_analyst_one_report", "")
    junior_analyst_two_report = state.get("junior_analyst_two_report", "")

    if not junior_analyst_one_report or not junior_analyst_two_report:
        logger.error("Missing junior analyst reports.")
        state["senior_analyst_report"] = "Cannot generate final report."
        return {"senior_analyst_report": state["senior_analyst_report"]}

    prompt = SENIOR_ANALYST_PROMPT.format(report1=junior_analyst_one_report, report2=junior_analyst_two_report)

    response = await anthropic_llm.ainvoke(prompt)
    # response = await gemini_llm.ainvoke(prompt) # FOR TESTING

    final_report = response.content if hasattr(response, "content") else str(response)

    query = ""
    for msg in state["messages"]:
        if isinstance(msg, HumanMessage):
            query = msg.content
            break

    query_safe = re.sub(r'[^\w\s-]', '', query)
    query_safe = query_safe.replace(" ", "_")[:50]
    file_name = f"Research_Report_{query_safe}.doc"
    file_path = os.path.join(tempfile.gettempdir(), file_name)

    doc = Document()

    # Parse markdown-like headings and format
    for line in final_report.splitlines():
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.strip():
            para = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
        

    doc.save(file_path)
    state["senior_analyst_report"] = final_report
    state["final_doc_path"] = file_path

    return state
